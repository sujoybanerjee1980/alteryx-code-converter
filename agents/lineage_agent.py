import os
import csv
import pandas as pd
from vertexai.generative_models import GenerativeModel as genai
from agents.utils_agent import ensure_folder
from agents.llm_agent import get_llm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import re
from docx.shared import Inches


def parse_markdown_table(md):
    lines = [ln.strip() for ln in md.split('\n') if ln.strip() and '|' in ln]
    if len(lines) < 2:
        return [], []
    headers = [h.strip() for h in lines[0].split('|') if h.strip()]
    data = []
    for ln in lines[2:]:
        row = [c.strip() for c in ln.split('|') if c.strip()]
        if row:
            data.append(row)
    return headers, data


def parse_testcase_blocks(text):
    """
    Returns a list of dicts:
    [{id, title, description, steps}], ... for each test case found in the text.
    """
    # Split by test case "Test Case ID:" (robust to extra lines/spaces)
    blocks = re.split(r"(?=Test Case ID:)", text)
    testcases = []
    for block in blocks:
        if not block.strip():
            continue
        tc = {}
        # Use regex to extract the fields
        id_match = re.search(r"Test Case ID:\s*(.+)", block)
        title_match = re.search(r"Test Case Title:\s*(.+)", block)
        desc_match = re.search(r"Test Case Description:([\s\S]*?)(Test Steps:|$)", block)
        steps_match = re.search(r"Test Steps:\s*([\s\S]+)", block)

        tc["id"] = id_match.group(1).strip() if id_match else ""
        tc["title"] = title_match.group(1).strip() if title_match else ""
        tc["description"] = desc_match.group(1).strip() if desc_match else ""
        if steps_match:
            # Clean up steps, like turn "- step" or "1. step" into bullet points
            steps_raw = steps_match.group(1).strip()
            steps = []
            for line in steps_raw.splitlines():
                s = line.lstrip("-0123456789. ").strip()
                if s:
                    steps.append(s)
            tc["steps"] = steps
        else:
            tc["steps"] = []
        testcases.append(tc)
    return testcases

        
def save_testcase_word_file_from_blocks(testcases, outpath, logo_path = None):
    doc = Document()
    if logo_path:
        doc.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()
    doc.add_heading('Test Cases', 0)
    for tc in testcases:
        doc.add_heading(f'Test Case ID: {tc["id"]}', level=2)
        doc.add_heading('Test Case Title:', level=3)
        doc.add_paragraph(tc["title"])
        doc.add_heading('Test Case Description:', level=4)
        doc.add_paragraph(tc["description"])
        doc.add_heading('Test Steps:', level=4)
        for step in tc["steps"]:
            doc.add_paragraph(step, style='List Bullet')
        doc.add_paragraph('-' * 40)
        doc.add_paragraph('')
    doc.save(outpath)

def save_excel(headers, rows, outpath):
    fixed_rows = []
    for row in rows:
        if len(row) > len(headers):
            fixed_rows.append(row[:len(headers)])
        elif len(row) < len(headers):
            fixed_rows.append(row + [None]*(len(headers)-len(row)))
        else:
            fixed_rows.append(row)
    df = pd.DataFrame(fixed_rows, columns=headers)
    df.to_excel(outpath, index=False)

def save_csv(headers, rows, outpath):
    with open(outpath, "w", encoding="utf-8", newline='') as fout:
        writer = csv.writer(fout)
        writer.writerow(headers)
        for row in rows:
            # ensure length match header (else csv will be broken)
            fixed_row = row if len(row) == len(headers) else (row + ['']*(len(headers)-len(row)))
            writer.writerow(fixed_row)

def generate_lineage_reports(source_logic_dict, target_logic_dict, model_name, ai_provider, output_folder, source, target):
    get_llm(ai_provider)
    if model_name == "gemini-1.5-flash":
        model = genai("gemini-2.0-flash-001")
    elif model_name == "gemini-2.0-flash-001":
        model = genai("gemini-1.5-flash")
    else:
        model = genai("gemini-1.5-flash")

    lineage_excel_paths = []
    testcase_word_paths = []
    all_avg_completion = {}
    lineage_reports_folder = ensure_folder(os.path.join(output_folder, "lineage_reports"))
    testcase_reports_folder = ensure_folder(os.path.join(output_folder, "testcase_reports"))

    file_list = list(source_logic_dict or {})
    n_files = len(file_list)

    for src_fname, src_logic in (source_logic_dict or {}).items():
        src_base = os.path.splitext(os.path.basename(src_fname))[0]
        tgt_fname, tgt_logic = None, ""

        for t_key, t_val in (target_logic_dict or {}).items():
            tgt_base = os.path.splitext(os.path.basename(t_key))[0]
            if tgt_base.startswith(src_base):
                tgt_fname, tgt_logic = t_key, t_val
                break

        # --- Lineage
        lineage_prompt = f"""Given the following source and target functional logic for an ETL process, 
            create a Deatiled business lineage mapping as a markdown table:
            use these columns only
            S.No | Business Functional Logic | Source Functional Logic Description - {source.capitalize()} | Target Functional Logic Description - {target.capitalize()}| Transfer Mapping Logic | Completion in % | Accuracy | Validation Notes/Comments
            
            Important instructions:
            Do not leave(or provide) any column as blank or as N/A, do not provide extra columns data(Do not add extra columns.).                                                                                                               
            For completion %, provide an estimated percentage (eg./ 100%, 80%, etc.) based on how closely source and target logic match.
            Decompose the overall process into as many detailed, meaningful steps as possible.
            For accuracy use 'Very High'(if completion > 95%), 'High'(if completion > 90%), 'Medium'(if completion > 80%), 'Low'(if completion <80) depending on confidence in mapping, also compare completion with accuracy.
            Return only the table, not any prose.


            Source Functional Logic: {src_logic}
            Target Functional Logic: {tgt_logic}

            
        """
        # In Validation Notes/Comments column, if there is no equivalent code available in source and target functional logic then provide that line number of no equivalent line code based on both functional logics mapping.

        
        lineage_response = model.generate_content(lineage_prompt)
        headers, rows = parse_markdown_table(lineage_response.text.strip())
        lineage_excel = os.path.join(lineage_reports_folder, f"{src_base}_lineage.xlsx")
        save_excel(headers, rows, lineage_excel)
        lineage_excel_paths.append(lineage_excel)

        try:
            df = pd.read_excel(lineage_excel)
            col = [c for c in df.columns if "Completion" in c and "%" in c]
            col = col[0] if col else None
            if col:
                nums = pd.to_numeric(df[col].astype(str).str.replace('%', ''), errors = 'coerce')
                avg = nums.mean()
                all_avg_completion[os.path.basename(lineage_excel)] = round(avg, 2) if not pd.isna(avg) else ""
            else:
                all_avg_completion[os.path.basename(lineage_excel)] = ""

        except Exception as ex:
            all_avg_completion[os.path.basename(lineage_excel)] = ""

        testcase_prompt = f"""Given the following target functional logics, 
            generate the set of test cases for the target functional logic(Not in table form, but as word document layout).
            
            For each test case, output with these parts and clear markdown-like headings:
            -Test Case ID
            -Test Case Title
            -Test Case Description
            -Test Steps(as a list or paragraphs)

            Expand from the functional mapping: include scenarios for valid data, invalid data and corner cases.
            Use clear, business-focused language.

            Format:
            For each test case, preasent as:

            Test Case ID: TC-1
            Test Case Title: <short title>
            Test Case Description:
            <full description>
            Test Steps:
            -Step 1: ...
            -Step 2: ...

            (add a horizontal line or spacing between test cases)

            
            Target Logic: {tgt_logic}
        """

        testcase_response = model.generate_content(testcase_prompt)
        testcases = parse_testcase_blocks(testcase_response.text.strip())
        testcase_word = os.path.join(testcase_reports_folder, f"{src_base}_testcase.docx")
        logo_path = 'images/tcs-logo.png'
        save_testcase_word_file_from_blocks(testcases, testcase_word, logo_path)
        testcase_word_paths.append(testcase_word)

    return lineage_excel_paths, testcase_word_paths, n_files, all_avg_completion
